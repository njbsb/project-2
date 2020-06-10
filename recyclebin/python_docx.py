# from docx import Document
# from docx.shared import Inches

# document = Document()

# document.add_heading('Document Title', 0)

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )

# document.add_picture('monty-truth.png', width=Inches(1.25))

# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

# document.save('demo.docx')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
document = Document()
header = document.sections[0].header
htable = header.add_table(1, 2, Inches(6))
htab_cells = htable.rows[0].cells
ht0 = htab_cells[0].add_paragraph()
kh = ht0.add_run()
kh.add_picture('logo.jpg', width=Inches(1))
kh.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# ht1 = htab_cells[1].add_paragraph('put your header text here')
# ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
document.save('yourdoc.docx')
